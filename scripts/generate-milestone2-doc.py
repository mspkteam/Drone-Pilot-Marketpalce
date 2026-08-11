from docx import Document
from docx.shared import Pt

doc = Document()

p = doc.add_paragraph()
run = p.add_run("REMOTE AIR SERVICE")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
run = p.add_run("Milestone 2 Progress Update")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
run = p.add_run(
    "Admin, Moderator, and Super Admin platform operations — job approval, "
    "fleet management, disputes, commissions, CMS, configuration, and staff permissions"
)
run.italic = True

doc.add_paragraph()
doc.add_paragraph("Hi,")
doc.add_paragraph(
    "Please find below a complete progress update for Milestone 2 (Admin & Moderator), "
    "covering the staff operations platform for Super Admins, Admins, and Moderators. "
    "All listed features use real, persisted data in the Neon database — not placeholder or mock content."
)

doc.add_heading("Live Environments", level=2)
table = doc.add_table(rows=2, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Surface"
table.rows[0].cells[1].text = "URL"
table.rows[1].cells[0].text = "Main Platform — Production"
table.rows[1].cells[1].text = "https://drone-pilot-marketpalce.vercel.app"

doc.add_paragraph(
    "Staff dashboards are accessed after login. Use the demo credentials in the section below "
    "to review each role."
)

doc.add_heading("Milestone 2 — Scope Completion", level=2)
doc.add_paragraph(
    "The following Week 2 / Milestone 2 deliverables are complete and available for review:"
)

items = [
    ("2.1 Job Approval Queue", "Admin reviews submitted client jobs and approves or rejects them. Approved jobs become visible to pilots based on membership grade timing."),
    ("2.5 Fleet & Personnel", "Real roster of pilots, clients, and staff from the database — no placeholder entries."),
    ("2.6 Dispute Centre", "Full dispute resolution workflow — view cases, investigate, add entries, review, and resolve."),
    ("2.7 Commissions Ledger", "Live commission tracking on completed jobs at the platform default rate of 15% (per approved business rules)."),
    ("2.8 Moderator Permissions", "Control what each Moderator can see and do — enforced in navigation, page access, and server actions."),
    ("2.9 CMS Articles & Resources", "Create, edit, and publish help articles and downloadable resources; public resources page; direct media uploads for featured images and files."),
    ("2.10 Platform Configuration", "Fees, commission rules, custom pilot rates, email template previews, integration status, and security settings — persisted to the database."),
    ("2.11 Certificates Engine", "Manage certificate templates, live preview, and issue certificates to pilots."),
    ("2.12 Badges & Wings", "Define achievement wings/badges, award conditions, wing images, and manual assignment with saved notes."),
    ("2.13 Uniform Shop Admin", "Manage shop products, variants, and orders — including full order history and status updates."),
    ("2.14 Subscriptions Admin", "Manage subscription tier plans and pilot enrolment (billing integration planned for a later milestone)."),
    ("2.16 Action-Level Permissions", "Role-based access enforced on admin actions — Super Admin, Admin, and Moderator maps with API-level guards."),
]
for title, desc in items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(title + " — ")
    r.bold = True
    p.add_run(desc)

p = doc.add_paragraph(
    "Goal achieved: Platform operators can fully manage and control the marketplace with real, persisted data."
)
p.runs[0].italic = True

doc.add_heading("Staff Demo Logins", level=2)
doc.add_paragraph("Password for all staff demo accounts: Demo123!")
login_table = doc.add_table(rows=4, cols=2)
login_table.style = "Table Grid"
login_table.rows[0].cells[0].text = "Role"
login_table.rows[0].cells[1].text = "Email"
logins = [
    ("Super Admin", "admin@dronepilot.local"),
    ("Admin", "ops@dronepilot.local"),
    ("Moderator", "moderator@dronepilot.local"),
]
for i, (role, email) in enumerate(logins, 1):
    login_table.rows[i].cells[0].text = role
    login_table.rows[i].cells[1].text = email

doc.add_paragraph("Login route: /login → redirects to /dashboard/admin after authentication.")

doc.add_heading("Milestone 2 — Admin & Operations Dashboard", level=2)
doc.add_paragraph(
    "The staff workspace is available at /dashboard/admin with role-appropriate navigation. "
    "Moderators see a filtered sidebar based on their permission preset (Full, Limited, or Custom). "
    "Super Admins have unrestricted access including Configuration, Regions, and Permissions."
)

sections = [
    ("Operations Dashboard", "/dashboard/admin", [
        "Platform growth and mission volume metrics",
        "Live system integrity indicators",
        "Interactive charts and operational status",
    ]),
    ("Job Approval Queue", "/dashboard/admin/jobs", [
        "Review pending client projects",
        "Approve or reject with persisted status updates",
        "Approved jobs unlock for pilot visibility per grade rules",
    ]),
    ("Fleet & Personnel", "/dashboard/admin/users", [
        "Unified roster of pilots, clients, and management users",
        "Role labels, regions, wings count, and status",
        "Super Admin can create and manage Admin/Moderator accounts",
    ]),
    ("Verifications", "/dashboard/admin/verifications", [
        "Review pilot license, insurance, and identity documents",
        "Approve or reject with audit trail",
    ]),
    ("Dispute Centre", "/dashboard/admin/disputes", [
        "Dispute list and detailed case views",
        "Staff entries, review, and resolution workflow",
        "Commission impact reflected on resolution",
    ]),
    ("Commissions & Payments", "/dashboard/admin/payments", [
        "Live commission ledger at 15% default",
        "Per-pilot custom commission overrides (Super Admin Configuration)",
        "Payment records linked to completed bookings",
    ]),
    ("Reports & Analytics", "/dashboard/admin/reports", [
        "Revenue, mission, and segmentation analytics from live data",
        "Interactive charts and filters",
    ]),
    ("Subscriptions", "/dashboard/admin/subscriptions", [
        "A-1 through A-6 tier plan management",
        "Pricing aligned with membership business rules",
        "Pilot enrolment visibility",
    ]),
    ("Certificates", "/dashboard/admin/certificates", [
        "Template library with live certificate preview",
        "Issue and track pilot certificates",
    ]),
    ("Badges & Wings", "/dashboard/admin/achievements", [
        "Six Figma-aligned wing definitions with image upload",
        "Award condition catalog for automated assignment rules",
        "Manual badge assignment with persisted notes",
    ]),
    ("Uniform Shop", "/dashboard/admin/shop", [
        "WooCommerce-style product editor",
        "Product variants and inventory",
        "Full orders list with status and payment updates",
    ]),
    ("CMS", "/dashboard/admin/cms", [
        "Articles and resources overview",
        "Create, edit, schedule, and publish content",
        "Featured image and file uploads",
        "Public resources integration",
    ]),
    ("Platform Configuration", "/dashboard/admin/settings", [
        "Default and grade-based commission display",
        "Custom pilot commission overrides (search, edit, save)",
        "Email template previews",
        "Integration and security settings (preview until external services connected)",
    ]),
    ("Staff Permissions", "/dashboard/admin/permissions", [
        "Super Admin only — assign module permissions per Moderator",
        "Permission change audit log",
    ]),
    ("Waitlist", "/dashboard/admin/waitlist", [
        "View pre-launch and marketing waitlist signups",
        "Role and region filters",
    ]),
    ("Support Chat", "/dashboard/admin/support", [
        "Real-time support conversations with pilots and clients",
    ]),
    ("Messages Tracking", "/dashboard/admin/messages", [
        "Read-only view of client–pilot marketplace conversations",
    ]),
]

for title, route, bullets in sections:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    r = p.add_run("Route: ")
    r.bold = True
    p.add_run(route)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Role Differences (Moderator vs Super Admin)", level=2)
for b in [
    "Moderators access the same /dashboard/admin surface with a filtered navigation menu.",
    "Configuration, Regions, and Permissions are Super Admin only.",
    "Each Moderator's module access (view, create, edit, delete) is stored in the database and enforced on every page and action.",
    "The Grade progress card has been removed from staff dashboards — it applies to pilots only.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Platform Behaviour", level=2)
for b in [
    "All Milestone 2 admin flows use the live API and Neon database.",
    "Default platform commission is 15% on completed bookings (aligned with approved RAS business rules; early planning docs referenced 10%).",
    "Super Admins can set per-pilot commission overrides from Configuration → Custom Pilot Rates.",
    "Stripe payment processing, production email delivery, and external integration connections are planned for later milestones.",
    "Security toggles on the Configuration page are preview UI until authentication integrations are activated.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Quality Assurance — Next Phase", level=2)
doc.add_paragraph(
    "Milestone 2 development and structure are complete. Formal end-to-end testing and QA "
    "are scheduled as Phase 4 of the project plan. The full platform will be tested holistically, "
    "and submission for final approval will follow QA team sign-off. Any items noted during "
    "informal review now will be captured and addressed during the dedicated QA phase."
)

doc.add_heading("What We Need From You", level=2)
for b in [
    "Review of the Milestone 2 staff dashboards for Super Admin, Admin, and Moderator roles.",
    "Confirmation that the 15% commission rate (vs. 10% in early planning docs) is correct per business rules.",
    "Feedback on any Figma alignment gaps before we proceed to Milestone 3 (Pilot).",
    "Approval to proceed to Milestone 3 — Pilot marketplace, proposals, contracts, and deliverables.",
]:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Summary", level=2)
doc.add_paragraph("Milestone 2 (Admin & Moderator) delivery is now complete.")
doc.add_paragraph(
    "Super Admins, Admins, and Moderators can fully operate the marketplace: approving jobs, "
    "managing fleet and verifications, resolving disputes, tracking commissions, publishing CMS content, "
    "configuring platform fees and custom pilot rates, and controlling staff permissions — all backed by "
    "real persisted data on the production environment."
)
doc.add_paragraph(
    "The platform is live at https://drone-pilot-marketpalce.vercel.app. "
    "Please log in with the staff demo credentials above and share any feedback or revisions "
    "you would like addressed before we begin Milestone 3."
)

doc.add_paragraph()
doc.add_paragraph("Best regards,")

out = r"d:\Desktop\Remote_Air_Service_Milestone_2_Progress_Update.docx"
doc.save(out)
print("Saved:", out)
